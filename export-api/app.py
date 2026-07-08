#!/usr/bin/env python3
from flask import Flask, request, jsonify, send_file
from datetime import datetime
from collections import defaultdict
import requests, urllib3, io, os
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document as WordDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
app = Flask(__name__)

ES_HOST = "https://localhost:9200"
ES_USER = "elastic"
ES_PASS = "<ES_PASSWORD>"
ES_INDEX = "safesquid-extended-*"

# Only genuinely policy-violating categories
VIOLATION_CATEGORIES = [
    "Pornography/Sexually Explicit",
    "Malware",
    "Phishing",
    "Botnet",
    "Gambling",
    "Hacking",
    "Spyware",
    "Adware",
]

# Filter names that indicate a security/policy block
VIOLATION_FILTERS = ["dnsbl","security","security-restrictions","sqscan","dlp","imgfilter"]

# Filtering reasons that indicate a real violation
VIOLATION_REASONS = [
    "Suspicious Domain",
    "EICAR.TestFile",
    "score",
    "access policy restrictions",
    "We forbid all the users from uploading",
    "Unauthorized Connect Port",
]

def is_blocked(status, filter_name):
    s = (status or "").strip()
    fn = (filter_name or "").strip().lower()
    if s in ["451","403"]:
        return True
    if fn == "dnsbl":
        return True
    return False

def is_violation(categories, filter_name, filtering_reason):
    cats = (categories or "").strip()
    fn   = (filter_name or "").strip().lower()
    fr   = (filtering_reason or "").strip()
    for vc in VIOLATION_CATEGORIES:
        if vc.lower() in cats.lower(): return True
    for vf in VIOLATION_FILTERS:
        if vf.lower() == fn: return True
    for vr in VIOLATION_REASONS:
        if vr.lower() in fr.lower(): return True
    return False

def parse_access_type(profiles):
    p = (profiles or "").upper().replace('"','').strip()
    if "READ ONLY" in p or "MINIMAL" in p: return "Read-Only"
    if "FULL CONTENT" in p: return "Full Access"
    return p if p and p != "-" else "-"

def clean(val):
    return val.replace('"','').strip() if isinstance(val, str) else "-"

def fmt_ts(ts):
    try:
        dt = datetime.fromisoformat(ts.replace("Z","+00:00"))
        return dt.strftime("%d %b %Y  %H:%M:%S")
    except: return ts or "-"

def es_fetch(username, client_ip, from_dt, to_dt, violations_only=False):
    must = []
    if username:  must.append({"wildcard":{"username.keyword":f"*{username.lower()}*"}})
    if client_ip: must.append({"wildcard":{"client_ip.keyword":f"*{client_ip}*"}})
    rng = {}
    if from_dt: rng["gte"] = from_dt
    if to_dt:   rng["lte"] = to_dt
    if rng: must.append({"range":{"@timestamp":rng}})

    if violations_only:
        should = []
        for vc in VIOLATION_CATEGORIES:
            should.append({"wildcard":{"categories.keyword":f"*{vc}*"}})
        for vf in VIOLATION_FILTERS:
            should.append({"term":{"filter_name.keyword":vf}})
        for vr in VIOLATION_REASONS:
            should.append({"wildcard":{"filtering_reason.keyword":f"*{vr}*"}})
        must.append({"bool":{"should":should,"minimum_should_match":1}})

    q = {
        "query": {"bool":{"must":must}} if must else {"match_all":{}},
        "sort": [{"@timestamp":{"order":"desc"}}],
        "_source": ["@timestamp","username","client_ip","url","profiles","status",
                    "method","categories","user_groups","filter_name",
                    "filtering_reason","request_host"],
        "size": 2000
    }
    r = requests.post(f"{ES_HOST}/{ES_INDEX}/_search?scroll=2m",
        auth=(ES_USER,ES_PASS), verify=False,
        headers={"Content-Type":"application/json"}, json=q, timeout=30)
    r.raise_for_status()
    data = r.json()
    hits = data["hits"]["hits"]
    total = data["hits"]["total"]["value"]
    scroll_id = data.get("_scroll_id")
    MAX_EXPORT_RECORDS = 20000
    while scroll_id and len(hits) < total and len(hits) < MAX_EXPORT_RECORDS:
        sr = requests.post(f"{ES_HOST}/_search/scroll",
            auth=(ES_USER,ES_PASS), verify=False,
            headers={"Content-Type":"application/json"},
            json={"scroll":"2m","scroll_id":scroll_id}, timeout=30)
        batch = sr.json()["hits"]["hits"]
        if not batch: break
        hits.extend(batch)
        scroll_id = sr.json().get("_scroll_id")
    return hits, total

def to_row(src, vio=False):
    status = clean(src.get("status","-"))
    fname  = clean(src.get("filter_name","-"))
    row = {
        "timestamp":   fmt_ts(src.get("@timestamp","")),
        "username":    clean(src.get("username","-")),
        "client_ip":   clean(src.get("client_ip","-")),
        "url":         clean(src.get("url","-")),
        "access_type": parse_access_type(src.get("profiles","")),
        "status":      status,
        "method":      clean(src.get("method","-")),
        "categories":  clean(src.get("categories","-")),
        "user_groups": clean(src.get("user_groups","-")),
        "blocked":     is_blocked(status, fname),
    }
    if vio:
        row["filter_name"]      = clean(src.get("filter_name","-"))
        row["filtering_reason"] = clean(src.get("filtering_reason","-"))
        row["request_host"]     = clean(src.get("request_host","-"))
    return row


def es_aggregate_user_ips():
    """Use ES terms aggregation to get TRUE complete list of users and their IPs."""
    q = {
        "size": 0,
        "aggs": {
            "by_user": {
                "terms": {"field": "username.keyword", "size": 200},
                "aggs": {
                    "by_ip": {"terms": {"field": "client_ip.keyword", "size": 50}}
                }
            }
        }
    }
    r = requests.post(f"{ES_HOST}/{ES_INDEX}/_search",
        auth=(ES_USER,ES_PASS), verify=False,
        headers={"Content-Type":"application/json"}, json=q, timeout=30)
    r.raise_for_status()
    data = r.json()
    result = {}
    for bucket in data.get("aggregations",{}).get("by_user",{}).get("buckets",[]):
        uname = clean(bucket["key"])
        ips = [clean(ipb["key"]) for ipb in bucket.get("by_ip",{}).get("buckets",[])]
        result[uname] = sorted(ips)
    return result

def es_aggregate_users_summary():
    """Use ES aggregation to get TRUE complete user list with counts and last seen."""
    q = {
        "size": 0,
        "aggs": {
            "by_user": {
                "terms": {"field": "username.keyword", "size": 200},
                "aggs": {
                    "last_seen": {"max": {"field": "@timestamp"}},
                    "by_access": {"terms": {"field": "profiles.keyword", "size": 20}}
                }
            }
        }
    }
    r = requests.post(f"{ES_HOST}/{ES_INDEX}/_search",
        auth=(ES_USER,ES_PASS), verify=False,
        headers={"Content-Type":"application/json"}, json=q, timeout=30)
    r.raise_for_status()
    data = r.json()
    result = []
    for bucket in data.get("aggregations",{}).get("by_user",{}).get("buckets",[]):
        uname = clean(bucket["key"])
        total = bucket["doc_count"]
        last_seen_ms = bucket.get("last_seen",{}).get("value")
        last_seen = "-"
        if last_seen_ms:
            last_seen = fmt_ts(datetime.utcfromtimestamp(last_seen_ms/1000).isoformat()+"Z")
        full = 0
        readonly = 0
        for ab in bucket.get("by_access",{}).get("buckets",[]):
            at = parse_access_type(ab["key"])
            if at == "Full Access": full += ab["doc_count"]
            elif at == "Read-Only": readonly += ab["doc_count"]
        result.append({
            "username": uname, "total": total, "full": full,
            "readonly": readonly, "last_seen": last_seen, "violations": 0
        })
    return sorted(result, key=lambda x: x["total"], reverse=True)

@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"]  = "*"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    response.headers["Access-Control-Expose-Headers"] = "Content-Disposition, Content-Length, Content-Type"
    return response

@app.route("/api/report", methods=["POST","OPTIONS"])
def report():
    if request.method == "OPTIONS": return jsonify({}), 200
    b = request.get_json(force=True)
    username  = b.get("username","").strip()
    client_ip = b.get("client_ip","").strip()
    from_dt   = b.get("from_dt","").strip()
    to_dt     = b.get("to_dt","").strip()
    page      = max(1, int(b.get("page",1)))
    per_page  = int(b.get("per_page",50))
    try:
        hits, total = es_fetch(username, client_ip, from_dt, to_dt)
    except Exception as e:
        return jsonify({"error":str(e)}), 500
    rows  = [to_row(h["_source"]) for h in hits]
    start = (page-1)*per_page
    return jsonify({"total":total,"returned":len(rows),"page":page,
                    "per_page":per_page,"pages":max(1,-(-len(rows)//per_page)),
                    "rows":rows[start:start+per_page]})

@app.route("/api/violations", methods=["POST","OPTIONS"])
def violations():
    if request.method == "OPTIONS": return jsonify({}), 200
    b = request.get_json(force=True)
    username  = b.get("username","").strip()
    client_ip = b.get("client_ip","").strip()
    from_dt   = b.get("from_dt","").strip()
    to_dt     = b.get("to_dt","").strip()
    page      = max(1, int(b.get("page",1)))
    per_page  = int(b.get("per_page",50))
    try:
        hits, total = es_fetch(username, client_ip, from_dt, to_dt, violations_only=True)
    except Exception as e:
        return jsonify({"error":str(e)}), 500
    rows  = [to_row(h["_source"], vio=True) for h in hits]
    start = (page-1)*per_page
    return jsonify({"total":total,"returned":len(rows),"page":page,
                    "per_page":per_page,"pages":max(1,-(-len(rows)//per_page)),
                    "rows":rows[start:start+per_page]})

@app.route("/api/users", methods=["POST","OPTIONS"])
def users():
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        result = es_aggregate_users_summary()
    except Exception as e:
        return jsonify({"error":str(e)}), 500
    return jsonify({"users": result, "total": len(result)})
@app.route("/api/user_ips", methods=["POST","OPTIONS"])
def user_ips():
    if request.method == "OPTIONS": return jsonify({}), 200
    try:
        result = es_aggregate_user_ips()
    except Exception as e:
        return jsonify({"error":str(e)}), 500
    return jsonify({"map": result})

def build_excel(rows, filters, vio=False):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Policy Violations" if vio else "Proxy Audit Report"
    thin   = Side(style="thin", color="CCCCCC")
    border = Border(left=thin,right=thin,top=thin,bottom=thin)
    NAVY,BLUE,WHITE = "1F3864","2E75B6","FFFFFF"
    cols = 11 if vio else 10
    last = chr(64+cols)
    ws.merge_cells(f"A1:{last}1")
    c = ws["A1"]
    c.value = "HDFC AMC  -  " + ("Policy Violations Report" if vio else "Proxy Audit Report")
    c.font  = Font(name="Calibri",bold=True,size=14,color=WHITE)
    c.fill  = PatternFill("solid",fgColor=NAVY)
    c.alignment = Alignment(horizontal="center",vertical="center")
    ws.row_dimensions[1].height = 28
    ws.merge_cells(f"A2:{last}2")
    meta = f"Generated: {datetime.now().strftime('%d %b %Y %H:%M:%S')}   |   Records: {len(rows)}"
    if filters: meta += "   |   " + ",  ".join(filters)
    c2 = ws["A2"]
    c2.value = meta
    c2.font  = Font(name="Calibri",size=9,color="444444")
    c2.fill  = PatternFill("solid",fgColor="EEF4FA")
    c2.alignment = Alignment(horizontal="left",vertical="center",indent=1)
    ws.row_dimensions[2].height = 16
    if vio:
        HEADERS = ["Date & Time","Username","Client IP","URL Accessed","Domain",
                   "Category","Access Type","Status","Method","Security Filter","Reason"]
        WIDTHS  = [20,16,14,50,25,28,14,10,10,18,40]
    else:
        HEADERS = ["Date & Time","Username","Client IP","URL Accessed",
                   "Access Type","HTTP Status","Method","Category","User Group","Blocked"]
        WIDTHS  = [22,18,16,60,14,12,10,30,25,12]
    for ci,(h,w) in enumerate(zip(HEADERS,WIDTHS),1):
        cell = ws.cell(row=3,column=ci,value=h)
        cell.font      = Font(name="Calibri",bold=True,size=10,color=WHITE)
        cell.fill      = PatternFill("solid",fgColor="C00000" if vio else BLUE)
        cell.alignment = Alignment(horizontal="center",vertical="center")
        cell.border    = border
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.row_dimensions[3].height = 18
    for ri,row in enumerate(rows,start=4):
        if row.get("blocked"):
            fill = PatternFill("solid",fgColor="FFCDD2")
        else:
            fill = PatternFill("solid",fgColor=("FFF5F5" if vio and ri%2==0 else
                                                 "F5F7FA" if not vio and ri%2==0 else WHITE))
        values = [row["timestamp"],row["username"],row["client_ip"],row["url"],
                  row.get("request_host","-"),row["categories"],row["access_type"],
                  row["status"],row["method"],row.get("filter_name","-"),
                  row.get("filtering_reason","-")] if vio else \
                 [row["timestamp"],row["username"],row["client_ip"],row["url"],
                  row["access_type"],row["status"],row["method"],
                  row["categories"],row["user_groups"],
                  ("BLOCKED" if row.get("blocked") else "")]
        for ci,val in enumerate(values,1):
            cell = ws.cell(row=ri,column=ci,value=val)
            cell.fill      = fill
            cell.font      = Font(name="Calibri",size=9)
            cell.alignment = Alignment(vertical="center",wrap_text=(ci==4))
            cell.border    = border
        ws.row_dimensions[ri].height = 14
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:{last}{3+len(rows)}"
    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    return buf

def build_pdf(rows, filters, vio=False):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            rightMargin=10*mm,leftMargin=10*mm,
                            topMargin=12*mm,bottomMargin=12*mm)
    BLUE  = colors.HexColor("#2E75B6")
    RED   = colors.HexColor("#C00000")
    LGRAY = colors.HexColor("#F5F7FA")
    LRED  = colors.HexColor("#FFF5F5")
    GRAY  = colors.HexColor("#CCCCCC")
    cell_s = ParagraphStyle("c",fontSize=7,fontName="Helvetica",leading=9)
    meta_s = ParagraphStyle("m",fontSize=8,fontName="Helvetica")
    hdr_s  = ParagraphStyle("h",fontSize=8,textColor=colors.white,
                             fontName="Helvetica-Bold",alignment=TA_CENTER)
    title_text = "HDFC AMC  -  " + ("Policy Violations Report" if vio else "Proxy Audit Report")
    meta_text  = f"Generated: {datetime.now().strftime('%d %b %Y %H:%M:%S')}   |   Records: {len(rows)}"
    if filters: meta_text += "   |   " + ",  ".join(filters)
    if vio:
        headers = ["Date & Time","Username","Client IP","URL","Category","Status","Method","Filter","Reason"]
        col_w   = [42*mm,25*mm,22*mm,52*mm,35*mm,13*mm,13*mm,20*mm,48*mm]
        def gv(r): return [r["timestamp"],r["username"],r["client_ip"],
                            Paragraph(r["url"][:80],cell_s),r["categories"],
                            r["status"],r["method"],r.get("filter_name","-"),
                            Paragraph(r.get("filtering_reason","-")[:100],cell_s)]
        row_bgs = [colors.white, LRED]
    else:
        headers = ["Date & Time","Username","Client IP","URL","Access Type","Status","Method","Category","Blocked"]
        col_w   = [38*mm,22*mm,20*mm,62*mm,22*mm,12*mm,12*mm,32*mm,16*mm]
        def gv(r): return [r["timestamp"],r["username"],r["client_ip"],
                            Paragraph(r["url"][:80],cell_s),r["access_type"],
                            r["status"],r["method"],r["categories"],
                            ("BLOCKED" if r.get("blocked") else "")]
        row_bgs = [colors.white, LGRAY]
    data = [[Paragraph(h,hdr_s) for h in headers]]
    for r in rows: data.append(gv(r))
    tbl = Table(data,colWidths=col_w,repeatRows=1)
    style_cmds = [
        ("BACKGROUND",(0,0),(-1,0),RED if vio else BLUE),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),row_bgs),
        ("GRID",(0,0),(-1,-1),0.3,GRAY),
        ("FONTSIZE",(0,1),(-1,-1),7),
        ("FONTNAME",(0,1),(-1,-1),"Helvetica"),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),3),
        ("BOTTOMPADDING",(0,0),(-1,-1),3),
    ]
    BLOCKED_RED = colors.HexColor("#FFCDD2")
    for ridx,rrow in enumerate(rows, start=1):
        if rrow.get("blocked"):
            style_cmds.append(("BACKGROUND",(0,ridx),(-1,ridx),BLOCKED_RED))
    tbl.setStyle(TableStyle(style_cmds))
    hdr_color = colors.HexColor("#7F1D1D") if vio else colors.HexColor("#1F3864")
    story = [Paragraph(title_text,ParagraphStyle("t",fontSize=14,textColor=hdr_color,fontName="Helvetica-Bold")),
             Spacer(1,3*mm),Paragraph(meta_text,meta_s),Spacer(1,4*mm),tbl]
    doc.build(story); buf.seek(0)
    return buf

def set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
    tcPr.append(shd)

def build_word(rows, filters, vio=False):
    doc = WordDocument()
    for section in doc.sections:
        section.top_margin=Cm(1.5); section.bottom_margin=Cm(1.5)
        section.left_margin=Cm(1.5); section.right_margin=Cm(1.5)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HDFC AMC  -  " + ("Policy Violations Report" if vio else "Proxy Audit Report"))
    run.bold=True; run.font.size=Pt(16)
    run.font.color.rgb = RGBColor(0x7F,0x1D,0x1D) if vio else RGBColor(0x1F,0x38,0x64)
    meta_text = f"Generated: {datetime.now().strftime('%d %b %Y %H:%M:%S')}   |   Records: {len(rows)}"
    if filters: meta_text += "   |   " + ",  ".join(filters)
    meta = doc.add_paragraph(meta_text)
    meta.runs[0].font.size=Pt(8); meta.runs[0].font.color.rgb=RGBColor(0x44,0x44,0x44)
    doc.add_paragraph()
    if vio:
        headers = ["Date & Time","Username","Client IP","URL","Category","Status","Method","Security Filter","Reason"]
        col_w   = [Cm(3.2),Cm(2.5),Cm(2.2),Cm(5.5),Cm(3.5),Cm(1.4),Cm(1.4),Cm(2.5),Cm(4.5)]
        hdr_bg  = "C00000"
    else:
        headers = ["Date & Time","Username","Client IP","URL","Access Type","Status","Method","Category","User Group","Blocked"]
        col_w   = [Cm(3.2),Cm(2.2),Cm(2),Cm(6),Cm(2.2),Cm(1.3),Cm(1.3),Cm(2.7),Cm(2.2),Cm(1.6)]
        hdr_bg  = "2E75B6"
    tbl = doc.add_table(rows=1,cols=len(headers))
    tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,w in enumerate(col_w): tbl.columns[i].width=w
    hdr_row = tbl.rows[0]
    for i,h in enumerate(headers):
        cell=hdr_row.cells[i]; cell.text=h
        cell.paragraphs[0].runs[0].bold=True
        cell.paragraphs[0].runs[0].font.size=Pt(8)
        cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell,hdr_bg)
    for ri,row in enumerate(rows):
        values = [row["timestamp"],row["username"],row["client_ip"],
                  row["url"][:80],row["categories"],row["status"],row["method"],
                  row.get("filter_name","-"),row.get("filtering_reason","-")[:100]] if vio else \
                 [row["timestamp"],row["username"],row["client_ip"],
                  row["url"][:80],row["access_type"],row["status"],row["method"],
                  row["categories"],row["user_groups"],
                  ("BLOCKED" if row.get("blocked") else "")]
        tr=tbl.add_row()
        if row.get("blocked"):
            bg="FFCDD2"
        else:
            bg="FFF5F5" if vio and ri%2==0 else "F5F7FA" if not vio and ri%2==0 else "FFFFFF"
        for ci,val in enumerate(values):
            cell=tr.cells[ci]; cell.text=str(val)
            cell.paragraphs[0].runs[0].font.size=Pt(7.5)
            set_cell_bg(cell,bg)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

@app.route("/api/export", methods=["POST","OPTIONS"])
def export():
    if request.method == "OPTIONS": return jsonify({}), 200
    b         = request.get_json(force=True)
    username  = b.get("username","").strip()
    client_ip = b.get("client_ip","").strip()
    from_dt   = b.get("from_dt","").strip()
    to_dt     = b.get("to_dt","").strip()
    fmt       = b.get("format","excel").lower()
    vio       = b.get("violations_only",False)
    usernames = b.get("usernames",[])
    try:
        if usernames:
            all_hits=[]
            for u in usernames:
                h,_=es_fetch(u,client_ip,from_dt,to_dt,vio); all_hits.extend(h)
            hits=all_hits
        else:
            hits,_=es_fetch(username,client_ip,from_dt,to_dt,vio)
    except Exception as e:
        return jsonify({"error":str(e)}), 500
    rows = [to_row(h["_source"],vio=vio) for h in hits]
    filters=[]
    if username:  filters.append(f"Username: {username}")
    if client_ip: filters.append(f"IP: {client_ip}")
    if from_dt:   filters.append(f"From: {from_dt}")
    if to_dt:     filters.append(f"To: {to_dt}")
    if vio:       filters.append("Policy Violations Only")
    ts    = datetime.now().strftime("%Y%m%d_%H%M%S")
    rtype = "Violations" if vio else "Audit"
    if fmt == "pdf":
        buf=build_pdf(rows,filters,vio); fname=f"HDFC_{rtype}_{ts}.pdf"
        mime="application/pdf"
    elif fmt == "word":
        buf=build_word(rows,filters,vio); fname=f"HDFC_{rtype}_{ts}.docx"
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        buf=build_excel(rows,filters,vio); fname=f"HDFC_{rtype}_{ts}.xlsx"
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return send_file(buf,as_attachment=True,attachment_filename=fname,mimetype=mime)

@app.route("/")
def index(): return "HDFC AMC Proxy Audit API v2 - Running", 200


@app.route("/api/download", methods=["GET","OPTIONS"])
def download():
    if request.method == "OPTIONS": return jsonify({}), 200
    username  = request.args.get("username","").strip()
    client_ip = request.args.get("client_ip","").strip()
    from_dt   = request.args.get("from_dt","").strip()
    to_dt     = request.args.get("to_dt","").strip()
    fmt       = request.args.get("format","excel").lower()
    vio       = request.args.get("violations_only","false").lower() == "true"
    try:
        hits,_ = es_fetch(username, client_ip, from_dt, to_dt, vio)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    rows = [to_row(h["_source"], vio=vio) for h in hits]
    filters = []
    if username:  filters.append(f"Username: {username}")
    if client_ip: filters.append(f"IP: {client_ip}")
    if from_dt:   filters.append(f"From: {from_dt}")
    if to_dt:     filters.append(f"To: {to_dt}")
    if vio:       filters.append("Policy Violations Only")
    ts    = datetime.now().strftime("%Y%m%d_%H%M%S")
    rtype = "Violations" if vio else "Audit"
    if fmt == "pdf":
        buf=build_pdf(rows,filters,vio); fname=f"HDFC_{rtype}_{ts}.pdf"
        mime="application/pdf"
    elif fmt == "word":
        buf=build_word(rows,filters,vio); fname=f"HDFC_{rtype}_{ts}.docx"
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        buf=build_excel(rows,filters,vio); fname=f"HDFC_{rtype}_{ts}.xlsx"
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return send_file(buf, as_attachment=True, attachment_filename=fname, mimetype=mime)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001, debug=False)
